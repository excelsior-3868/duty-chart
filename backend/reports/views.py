# reports/views.py
import io
import datetime
import nepali_datetime

from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.response import Response
from django.utils import timezone
import datetime
from datetime import timedelta
from django.contrib.auth import get_user_model

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side

from duties.models import Duty, DutyChart
from .permissions import IsAdminOrSelf
from users.permissions import user_has_permission_slug, get_allowed_office_ids
import requests

User = get_user_model()


_translation_cache = {}

def translate_to_nepali(text):
    """Google Translate free API for simple names/titles with simple memory cache."""
    if not text or text == "-":
        return text
    
    # Return from cache if available
    if text in _translation_cache:
        return _translation_cache[text]
        
    try:
        url = "https://translate.googleapis.com/translate_a/single"
        params = {"client": "gtx", "sl": "en", "tl": "ne", "dt": "t", "q": text}
        response = requests.get(url, params=params, timeout=5)
        if response.status_code == 200:
            result = response.json()
            if result and result[0] and result[0][0]:
                translated = result[0][0][0]
                _translation_cache[translated] = translated # double save for consistency
                _translation_cache[text] = translated
                return translated
    except Exception as e:
        print(f"Translation error for '{text}': {e}")
        pass
    return text


def translate_to_nepali_batch(text_list):
    """
    Translates a list of strings in batches to minimize API calls.
    Uses \n as a delimiter for the gtx free API.
    """
    if not text_list:
        return {}
    
    # Filter out empty or already cached
    to_translate = [t for t in text_list if t and isinstance(t, str) and str(t) not in _translation_cache]
    if not to_translate:
        return {t: _translation_cache.get(str(t), str(t)) for t in text_list}
    
    # Process in chunks of 50 to stay safe with URL length
    chunk_size = 50
    for i in range(0, len(to_translate), chunk_size):
        chunk = to_translate[i:i + chunk_size]
        combined = "\n".join(chunk)
        
        try:
            url = "https://translate.googleapis.com/translate_a/single"
            params = {
                "client": "gtx",
                "sl": "en",
                "tl": "ne",
                "dt": "t",
                "q": combined
            }
            response = requests.get(url, params=params, timeout=10)
            if response.status_code == 200:
                result = response.json()
                if result and result[0]:
                    # result[0] is typically a list of [[translated_chunk, original_chunk, ...], ...]
                    translated_combined = "".join([part[0] for part in result[0] if part[0]])
                    translated_list = translated_combined.split("\n")
                    
                    # Map back to original strings
                    for orig, trans in zip(chunk, translated_list):
                        _translation_cache[str(orig)] = trans.strip()
        except Exception as e:
            print(f"Batch translation error: {e}")
            
    return {t: _translation_cache.get(str(t), str(t)) for t in text_list}


# ---------------------------
# Helper: Parse user_id[] or comma-separated
# ---------------------------
def _parse_user_ids(request):
    raw = request.GET.getlist("user_id[]") or request.GET.getlist("user_id")
    if not raw:
        s = request.GET.get("user_id")
        if s:
            raw = [s]

    ids = []
    for entry in raw:
        for part in str(entry).split(","):
            try:
                ids.append(int(part.strip()))
            except Exception:
                continue

    # dedupe, preserve order
    seen, out = set(), []
    for i in ids:
        if i not in seen:
            seen.add(i)
            out.append(i)
    return out


# ---------------------------
# Duty options (dropdown)
# ---------------------------
class DutyOptionsView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        qs = DutyChart.objects.all().order_by("effective_date")
        return Response([
            {
                "id": c.id,
                "name": c.name or f"{c.office.name} - {c.effective_date}",
                "effective_date": str(c.effective_date),
                "end_date": str(c.end_date) if c.end_date else str(c.effective_date),
                "office_id": c.office_id,
                "office_name": c.office.name if c.office else "Unknown",
            }
            for c in qs
        ])


# ---------------------------
# Preview JSON (unchanged behavior)
# ---------------------------
class DutyReportPreviewView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")
        if not (date_from and date_to):
            return Response({"groups": []})

        all_users = request.GET.get("all_users") == "1"
        user_ids = [] if all_users else _parse_user_ids(request)
        duty_id = request.GET.get("duty_id")
        schedule_id = request.GET.get("schedule_id")
        office_id = request.GET.get("office_id")

        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")

        qs = Duty.objects.select_related("user", "schedule", "office").filter(
            date__range=[date_from, date_to]
        )

        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        if user_ids:
            qs = qs.filter(user_id__in=user_ids)
        
        if duty_id:
            qs = qs.filter(duty_chart_id=duty_id)

        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)
        
        if office_id and office_id != "all":
            qs = qs.filter(office_id=office_id)

        qs = qs.order_by("user_id", "date", "schedule__start_time")

        groups = {}
        for d in qs:
            if not d.user:
                continue
            uid = d.user.id
            groups.setdefault(uid, {
                "user_id": uid,
                "user_name": getattr(d.user, "full_name", str(d.user)),
                "employee_id": getattr(d.user, "employee_id", "-"),
                "office": d.office.name if d.office else "-",
                "rows": []
            })["rows"].append({
                "id": d.id,
                "date": str(d.date),
                "weekday": d.date.strftime("%A"),
                "schedule": getattr(d.schedule, "name", "-"),
                "start_time": str(d.schedule.start_time) if d.schedule else "-",
                "end_time": str(d.schedule.end_time) if d.schedule else "-",
                "is_completed": d.is_completed,
                "currently_available": d.currently_available,
            })

        return Response({"groups": list(groups.values())})


# ---------------------------
# DOCX EXPORT — अनुसूची-१
# FULL REPLACEMENT (FINAL)
# ---------------------------
class DutyReportFileView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        # -----------------------
        # PARAMS
        # -----------------------
        duty_id = request.GET.get("duty_id")
        all_users = request.GET.get("all_users") == "1"
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")

        # -----------------------
        # FULL CHART MODE
        # -----------------------
        if all_users and duty_id:
            try:
                chart = DutyChart.objects.get(id=duty_id)
                date_from = chart.effective_date
                date_to = chart.end_date or chart.effective_date
            except DutyChart.DoesNotExist:
                return Response({"error": "Invalid duty chart"}, status=400)

        # -----------------------
        # RANGE MODE VALIDATION
        # -----------------------
        if not (date_from and date_to):
            return Response({"error": "Missing dates"}, status=400)

        # -----------------------
        # QUERYSET
        # -----------------------
        qs = Duty.objects.select_related(
            "user", "schedule", "office"
        ).filter(
            date__range=[date_from, date_to]
        )

        if duty_id:
            qs = qs.filter(duty_chart_id=duty_id)

        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")
        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        qs = qs.order_by("date", "schedule__start_time")

        # -----------------------
        # EXCEL GENERATION
        # -----------------------
        fmt = request.GET.get("format", "docx")
        if fmt == "excel":
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Duty Chart"

            # Headers
            headers = [
                "SN", "Designation", "Name", "Phone", 
                "Work Description", "Target", "Timeline", "Remarks"
            ]
            ws.append(headers)

            # Styling Headers
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center")

            # Data
            sn = 1
            for d in qs:
                row = [
                    sn,
                    (d.user.position.alias or d.user.position.name) if d.user.position else "-",
                    getattr(d.user, "full_name", "-"),
                    getattr(d.user, "phone_number", "-"),
                    "", # Work Description
                    "", # Target
                    str(d.date),
                    ""  # Remarks
                ]
                ws.append(row)
                sn += 1
            
            # Auto-adjust column width (simple)
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter # Get the column name
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                ws.column_dimensions[column].width = adjusted_width

            bio = io.BytesIO()
            wb.save(bio)
            bio.seek(0)

            return FileResponse(
                bio,
                as_attachment=True,
                filename=f"Duty_Chart_{date_from}_{date_to}.xlsx",
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        # -----------------------
        # DOCX GENERATION
        # -----------------------
        doc = Document()
        doc.styles["Normal"].font.size = Pt(11)

        def center(text, bold=False):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = bold
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Header
        center("अनुसूची-१", True)
        center("(परिच्छेद - ३ को दफा ८ र १० सँग सम्बन्धित)")
        center("नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)")
        center("सिफ्ट ड्यूटीमा खटाउनु अघि भर्नु पर्ने बिवरण")

        # Convert dates to Nepali
        try:
            # Handle string or date objects
            d_from = datetime.datetime.strptime(str(date_from), "%Y-%m-%d") if isinstance(date_from, str) else date_from
            d_to = datetime.datetime.strptime(str(date_to), "%Y-%m-%d") if isinstance(date_to, str) else date_to
            
            nepali_from = nepali_datetime.date.from_datetime_date(d_from)
            nepali_to = nepali_datetime.date.from_datetime_date(d_to)
            nepali_period = f"{nepali_from} देखि {nepali_to} सम्म"
        except Exception as e:
            print(f"Failed to convert to Nepali date: {e}")
            nepali_period = f"{date_from} देखि {date_to} सम्म"

        # Unique schedules in this report for classification
        unique_schedules = []
        seen_schedules = set()
        for d in qs:
            if d.schedule and d.schedule.id not in seen_schedules:
                unique_schedules.append(d.schedule)
                seen_schedules.add(d.schedule.id)
        
        def format_time_nepali(t, crosses=False):
            t_str = t.strftime("%H:%M")
            return f"भोलिपल्ट {t_str}" if crosses else t_str

        classification_str = ""
        if unique_schedules:
            parts = []
            for s in unique_schedules:
                crosses = s.end_time < s.start_time
                start_t = s.start_time.strftime("%H:%M")
                end_t = format_time_nepali(s.end_time, crosses)
                parts.append(f"{s.name} ({start_t} — {end_t})")
            classification_str = ", ".join(parts)
        else:
            classification_str = "-"

        # Meta
        meta = doc.add_paragraph()
        meta.add_run("कार्यालयको नाम:- ").bold = True
        meta.add_run(qs.first().office.name if qs.exists() and qs.first().office else "-")

        meta.add_run("\nबिभाग/शाखाको नाम:- ").bold = True
        meta.add_run("\nमिति:- ").bold = True
        meta.add_run(nepali_period)

        meta.add_run("\nड्यूटीको बर्गिकरण:- ").bold = True
        meta.add_run(classification_str)

        doc.add_paragraph("\nकाममा खटाईएको बिवरण:-")

        # Table
        table = doc.add_table(rows=2, cols=8)
        table.style = "Table Grid"

        # Headers Setup with Merging
        # Vertical merges
        table.cell(0, 0).merge(table.cell(1, 0)) # सि.नं.
        table.cell(0, 4).merge(table.cell(1, 4)) # कामको बिवरण
        table.cell(0, 5).merge(table.cell(1, 5)) # लक्ष्य
        table.cell(0, 6).merge(table.cell(1, 6)) # समय सिमा
        table.cell(0, 7).merge(table.cell(1, 7)) # कैफियत

        # Horizontal merge
        table.cell(0, 1).merge(table.cell(0, 3)) # काममा खटाउनु पर्ने कर्मचारीहरुको बिवरण

        # Set Text and Alignment
        def set_cell_text(cell, text, bold=True):
            cell.text = text
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if bold:
                for run in p.runs:
                    run.bold = True

        set_cell_text(table.cell(0, 0), "सि.नं.")
        set_cell_text(table.cell(0, 1), "काममा खटाउनु पर्ने कर्मचारीहरुको विवरण")
        set_cell_text(table.cell(0, 4), "कामको विवरण")
        set_cell_text(table.cell(0, 5), "लक्ष्य")
        set_cell_text(table.cell(0, 6), "समय सिमा")
        
        # Timeline header with timings if single schedule
        timeline_header = "समय सिमा"
        is_single_schedule = unique_schedules and len(unique_schedules) == 1
        if is_single_schedule:
            s = unique_schedules[0]
            crosses = s.end_time < s.start_time
            start_t = s.start_time.strftime("%H:%M")
            end_t = format_time_nepali(s.end_time, crosses)
            timeline_header += f"\n({start_t} - {end_t})"
        set_cell_text(table.cell(0, 6), timeline_header)
        
        set_cell_text(table.cell(0, 7), "कैफियत")

        # Row 1 sub-headers
        set_cell_text(table.cell(1, 1), "पद")
        set_cell_text(table.cell(1, 2), "नाम")
        set_cell_text(table.cell(1, 3), "सम्पर्क नं.")

        # -----------------------
        # BATCH TRANSLATION
        # -----------------------
        unique_texts = set()
        for d in qs:
            if d.user:
                if d.user.full_name: unique_texts.add(d.user.full_name)
                if d.user.position:
                    if d.user.position.alias: unique_texts.add(d.user.position.alias)
                    else: unique_texts.add(d.user.position.name)
        
        translate_to_nepali_batch(list(unique_texts))

        # Helper for Nepali digits
        NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
        def nep(txt):
            return str(txt).translate(NEP_DIGITS)

        sn = 1
        for d in qs:
            cells = table.add_row().cells
            
            # SN
            cells[0].text = nep(f"{sn}.")
            
            # Position
            pos = d.user.position
            pos_nm = (pos.alias or pos.name) if pos else "-"
            cells[1].text = translate_to_nepali(pos_nm)
            
            # Name + Employee ID (in brackets)
            name_str = translate_to_nepali(getattr(d.user, "full_name", "-"))
            emp_id = getattr(d.user, "employee_id", "")
            if emp_id:
                name_str += f" ({nep(emp_id)})"
            cells[2].text = name_str
            
            # Phone
            cells[3].text = nep(getattr(d.user, "phone_number", "-"))
            
            # Work Description & Target
            cells[4].text = ""
            cells[5].text = ""

            # Date / Timeline in Nepali
            d_obj = d.date
            try:
                # Ensure d_obj is a date object
                if isinstance(d_obj, str):
                    d_obj = datetime.datetime.strptime(d_obj, "%Y-%m-%d").date()
                nepali_d = nepali_datetime.date.from_datetime_date(d_obj)
                date_str = nepali_d.strftime("%Y/%m/%d")
            except Exception:
                date_str = str(d_obj)
            
            date_str = nep(date_str)

            if d.schedule:
                if is_single_schedule:
                    cells[6].text = date_str
                else:
                    crosses = d.schedule.end_time < d.schedule.start_time
                    st = d.schedule.start_time.strftime("%H:%M")
                    et = format_time_nepali(d.schedule.end_time, crosses)
                    cells[6].text = f"{date_str}\n({nep(st)} - {nep(et)})"
            else:
                cells[6].text = date_str

            cells[7].text = ""

            # Align cells
            for i in range(8):
                p = cells[i].paragraphs[0]
                # Center most, left align name & work desc
                if i in [2, 4]:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                else:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            sn += 1

        # Footer
        doc.add_paragraph(
            "\nकम्पनीको सिफ्ट ड्यूटी निर्देशिका बमोजिम तपाईंहरुलाई माथि उल्लेखित "
            "समयसीमा भित्र कार्य सम्पन्न गर्ने गरी ड्यूटीमा खटाईएको छ | "
            "उक्त कार्य सम्पन्न गरे पश्चात् अनुसूची २ बमोजिम कार्य सम्पन्न गरेको "
            "प्रमाणित गराई पेश गर्नुहुन अनुरोध छ |"
        )

        doc.add_paragraph("\nकाममा खटाउने अधिकार प्राप्त पदाधिकारीको बिवरण:-")
        doc.add_paragraph("नाम:-")
        doc.add_paragraph("पद:-")
        doc.add_paragraph("दस्तखत:-")
        doc.add_paragraph("मिति:-")

        # Response
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        return FileResponse(
            bio,
            as_attachment=True,
            filename=f"Duty_Chart_{date_from}_{date_to}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


class DutyReportNewFileView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        duty_id = request.GET.get("duty_id")
        all_users = request.GET.get("all_users") == "1"
        user_ids = _parse_user_ids(request)
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")
        schedule_id = request.GET.get("schedule_id")
        office_id = request.GET.get("office_id")
        group_by_employee = request.GET.get("group_by_employee") == "true"

        chart = None
        if duty_id:
            try:
                chart = DutyChart.objects.get(id=duty_id)
                # If all_users or if dates weren't provided, use chart dates for filtering
                if all_users or not (date_from and date_to):
                    date_from = chart.effective_date
                    date_to = chart.end_date or chart.effective_date
            except DutyChart.DoesNotExist:
                return Response({"error": "Invalid duty chart"}, status=400)

        if not (date_from and date_to):
            return Response({"error": "Missing dates"}, status=400)

        qs = Duty.objects.select_related("user", "schedule", "office").filter(date__range=[date_from, date_to])
        if not all_users and user_ids:
            qs = qs.filter(user_id__in=user_ids)
        if duty_id:
            qs = qs.filter(duty_chart_id=duty_id)
        
        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)
            
        if office_id and office_id != "all":
            qs = qs.filter(office_id=office_id)
        
        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")
        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        qs = qs.order_by("date", "schedule__start_time")

        if group_by_employee:
            # Group duties by user
            grouped_data = {}
            for d in qs:
                uid = d.user_id
                if uid not in grouped_data:
                    grouped_data[uid] = {
                        'duty': d,
                        'dates': []
                    }
                grouped_data[uid]['dates'].append(d.date)
            
            # Reconstruct list of "representative" duties with consolidated dates
            # We'll use a trick: keep the 'duty' object but we'll override its date representation in the table loop
            processed_rows = []
            for uid, data in grouped_data.items():
                processed_rows.append({
                    'duty': data['duty'],
                    'date_str': ", ".join([dt.isoformat() for dt in sorted(data['dates'])])
                })
        else:
            processed_rows = [{'duty': d, 'date_str': d.date.isoformat()} for d in qs]

        doc = Document()
        
        # Moderate margins: Top/Bottom 1", Left/Right 0.75"
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        doc.styles["Normal"].font.size = Pt(11)

        def center(text, bold=False):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = bold
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        center("अनुसूची-२", True)
        center("(परिच्छेद - ३ को दफा ८,९ र १० सँग सम्बन्धित)")
        center("नेपाल दूरसंचार कम्पनी लिमिटेड (नेपाल टेलिकम)")
        center("सिफ्ट ड्यूटी सम्पन्न भए पश्चात भर्नु पर्ने बिवरण")
        meta = doc.add_paragraph()
        try:
            # If we have a chart, use its official period for the header as requested
            if chart:
                df = chart.effective_date
                dt = chart.end_date or chart.effective_date
            else:
                # Parse dates if they are strings
                if isinstance(date_from, str):
                    df = datetime.datetime.strptime(date_from, "%Y-%m-%d").date()
                else:
                    df = date_from
                    
                if isinstance(date_to, str):
                    dt = datetime.datetime.strptime(date_to, "%Y-%m-%d").date()
                else:
                    dt = date_to
                
            nepali_from = nepali_datetime.date.from_datetime_date(df)
            nepali_to = nepali_datetime.date.from_datetime_date(dt)
            
            # Format: YYYY/MM/DD in Nepali numerals
            nepali_period = f"{str(nepali_from).replace('-', '/')} देखि {str(nepali_to).replace('-', '/')} सम्म"
            NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
            nepali_period = nepali_period.translate(NEP_DIGITS)
        except Exception as e:
            print(f"Date conversion error: {e}")
            nepali_period = f"{date_from} देखि {date_to} सम्म"

        unique_schedules = []
        seen_schedules = set()
        for d in qs:
            if d.schedule and d.schedule.id not in seen_schedules:
                unique_schedules.append(d.schedule)
                seen_schedules.add(d.schedule.id)

        meta = doc.add_paragraph()
        meta.add_run("कार्यालयको नाम:- ").bold = True
        meta.add_run(qs.first().office.name if qs.exists() and qs.first().office else "-")
        meta.add_run("\nबिभाग/शाखाको नाम:- ").bold = True
        meta.add_run("\nमिति:- ").bold = True
        meta.add_run(nepali_period)

        # Duty Classification (बर्गिकरण)
        NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
        def nep(txt): return str(txt).translate(NEP_DIGITS)

        def format_time_nepali(t, crosses=False):
            t_str = t.strftime("%H:%M")
            return f"भोलिपल्ट {t_str}" if crosses else t_str

        parts = []
        for s in unique_schedules:
            crosses = s.end_time < s.start_time
            st = s.start_time.strftime("%H:%M")
            et = format_time_nepali(s.end_time, crosses)
            # Use Nepali digits for times as well
            parts.append(f"{s.name} ({nep(st)} - {nep(et)})")
        
        classification_str = ", ".join(parts) if parts else "-"
        meta.add_run("\nड्यूटीको बर्गिकरण:- ").bold = True
        meta.add_run(classification_str)

        doc.add_paragraph() # Spacer
        p_desc = doc.add_paragraph()
        run_desc = p_desc.add_run("सम्पन्न भएको कामको बिवरण:-")
        run_desc.bold = True

        table = doc.add_table(rows=2, cols=8)
        table.style = "Table Grid"

        # Adjust Column Widths
        table.columns[0].width = Inches(0.25) # S.N.
        table.columns[1].width = Inches(0.5)  # Position
        table.columns[2].width = Inches(2.45) # Name
        table.columns[3].width = Inches(0.8)  # Contact
        table.columns[4].width = Inches(0.8)  # Work
        table.columns[5].width = Inches(0.6)  # Target
        table.columns[6].width = Inches(1.3)  # Achievement/Time
        table.columns[7].width = Inches(0.6)  # Remarks

        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 4).merge(table.cell(1, 4))
        table.cell(0, 5).merge(table.cell(1, 5))
        table.cell(0, 6).merge(table.cell(1, 6))
        table.cell(0, 7).merge(table.cell(1, 7))
        table.cell(0, 1).merge(table.cell(0, 3))

        def set_cell_text(cell, text, bold=True):
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            if bold:
                for run in p.runs: run.bold = True

        set_cell_text(table.cell(0, 0), "सि.नं.")
        set_cell_text(table.cell(0, 1), "कर्मचारीहरुले सम्पादन गरेको कामको बिवरण")
        set_cell_text(table.cell(0, 4), "कामको विवरण")
        set_cell_text(table.cell(0, 5), "लक्ष्य")
        set_cell_text(table.cell(0, 6), "उपलब्धि")
        set_cell_text(table.cell(0, 7), "काम सम्पादन नभएमा सो को कारण")

        timeline_header = "काम सम्पादन गर्न लागेको समय"
        is_single_schedule = unique_schedules and len(unique_schedules) == 1
        if is_single_schedule:
            s = unique_schedules[0]
            st = s.start_time.strftime("%H:%M")
            et = s.end_time.strftime("%H:%M")
            timeline_header += f"\n({nep(st)} - {nep(et)})"
        set_cell_text(table.cell(1, 6), timeline_header) # achievement subheader essentially

        set_cell_text(table.cell(1, 1), "पद")
        set_cell_text(table.cell(1, 2), "नाम")
        set_cell_text(table.cell(1, 3), "सम्पर्क नं.")

        # -----------------------
        # BATCH TRANSLATION
        # -----------------------
        unique_texts = set()
        for d in qs:
            if d.user:
                if d.user.full_name: unique_texts.add(d.user.full_name)
                if d.user.position:
                    if d.user.position.alias: unique_texts.add(d.user.position.alias)
                    else: unique_texts.add(d.user.position.name)
        
        translate_to_nepali_batch(list(unique_texts))

        NEP_DIGITS = str.maketrans("0123456789", "०१२३४५६७८९")
        def nep(txt): return str(txt).translate(NEP_DIGITS)

        sn = 1
        for row_item in processed_rows:
            d = row_item['duty']
            date_input = row_item['date_str']

            cells = table.add_row().cells
            cells[0].text = nep(f"{sn}.")
            
            # Position (पद) - Use pre-translated alias/name
            pos = d.user.position
            pos_nm = (pos.alias or pos.name) if pos else "-"
            cells[1].text = translate_to_nepali(pos_nm)
            
            # Name - Use pre-translated full_name
            raw_name = getattr(d.user, "full_name", "-")
            name_str = translate_to_nepali(raw_name)
            emp_id = getattr(d.user, "employee_id", "")
            if emp_id: name_str += f" ({nep(emp_id)})"
            
            cells[2].text = name_str
            cells[3].text = nep(getattr(d.user, "phone_number", "-"))
            cells[4].text = ""
            cells[5].text = ""
            
            # Date / Timeline in Nepali
            date_parts = [dp.strip() for dp in date_input.split(',')]
            nepali_dates = []
            for d_str in date_parts:
                try:
                    d_obj = datetime.datetime.strptime(d_str, "%Y-%m-%d").date()
                    nepali_d = nepali_datetime.date.from_datetime_date(d_obj)
                    nepali_dates.append(str(nepali_d).replace("-", "/"))
                except:
                    nepali_dates.append(d_str.replace("-", "/"))
            
            cells[6].text = nep(", ".join(nepali_dates))
            cells[7].text = ""

            for i in range(8):
                cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if i in [2, 4] else WD_PARAGRAPH_ALIGNMENT.CENTER
            sn += 1

        
        doc.add_paragraph("")
        doc.add_paragraph("काम सम्पादन भएको प्रमाणित गर्ने पदाधिकारीको विवरण :-")
        
        # Create a 2-column table for signatures
        sig_table = doc.add_table(rows=5, cols=2)
        sig_table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Column 1: पेश गर्ने
        sig_table.cell(0, 0).paragraphs[0].add_run("पेश गर्ने:").underline = True
        sig_table.cell(1, 0).text = "नाम :-"
        sig_table.cell(2, 0).text = "पद :-"
        sig_table.cell(3, 0).text = "दस्तखत:-"
        sig_table.cell(4, 0).text = "मिति :-"

        # Column 2: प्रमाणित गर्ने
        sig_table.cell(0, 1).paragraphs[0].add_run("प्रमाणित गर्ने:").underline = True
        sig_table.cell(1, 1).text = "नाम :-"
        sig_table.cell(2, 1).text = "पद :-"
        sig_table.cell(3, 1).text = "दस्तखत:-"
        sig_table.cell(4, 1).text = "मिति :-"

        # Set left indentation for signature fields within columns
        for row in sig_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.left_indent = Inches(0.5)

        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return FileResponse(bio, as_attachment=True, filename=f"Duty_Report_New_{date_from}_{date_to}.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

class SummaryReportView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        office_id = request.GET.get("office_id")
        date_from = request.GET.get("date_from")
        date_to = request.GET.get("date_to")
        user_id = request.GET.get("user_id")
        user_ids = request.GET.get("user_ids")
        schedule_id = request.GET.get("schedule_id")

        if not (date_from and date_to):
            return Response({"error": "Date range is required"}, status=400)

        qs = Duty.objects.select_related(
            "user", 
            "user__office", 
            "user__office__directorate", 
            "user__office__ac_office", 
            "user__office__cc_office",
            "schedule", 
            "office", 
            "duty_chart"
        ).filter(
            date__range=[date_from, date_to]
        )

        # Permission check
        can_see_any_office = request.user.is_staff or user_has_permission_slug(request.user, "duties.create_any_office_chart")
        if not can_see_any_office:
            allowed_offices = get_allowed_office_ids(request.user)
            qs = qs.filter(office_id__in=allowed_offices)

        if office_id and office_id != "all":
            qs = qs.filter(office_id=office_id)
        
        if user_id:
            qs = qs.filter(user_id=user_id)
        
        if user_ids:
            try:
                uid_list = [int(u) for u in user_ids.split(",") if u.strip().isdigit()]
                if uid_list:
                    qs = qs.filter(user_id__in=uid_list)
            except:
                pass

        if schedule_id and schedule_id != "all":
            qs = qs.filter(schedule_id=schedule_id)

        # Aggregate data
        summary = {}
        for d in qs:
            if not d.user: continue
            uid = d.user.id
            if uid not in summary:
                summary[uid] = {
                    "user_id": uid,
                    "full_name": d.user.full_name,
                    "employee_id": d.user.employee_id,
                    "office_name": d.office.name if d.office else "-",
                    "total_duties": 0,
                    "total_hours": 0.0,
                    "chart_breakdown": {}, # { chart_name: { shift_name: count } }
                    "dates": []
                }
            
            summary[uid]["total_duties"] += 1
            summary[uid]["dates"].append({
                "date": str(d.date),
                "chart": d.duty_chart.name if d.duty_chart and d.duty_chart.name else "Other/Manual",
                "shift": d.schedule.name if d.schedule else "No Shift",
                "day": d.date.strftime('%A')
            })
            
            chart_name = d.duty_chart.name if d.duty_chart and d.duty_chart.name else "Other/Manual"
            if chart_name not in summary[uid]["chart_breakdown"]:
                summary[uid]["chart_breakdown"][chart_name] = {
                    "total_duties": 0,
                    "total_hours": 0.0,
                    "shifts": {}
                }
            
            summary[uid]["chart_breakdown"][chart_name]["total_duties"] += 1
            
            # Calculate hours and count shift
            if d.schedule:
                s = d.schedule
                start = datetime.datetime.combine(datetime.date.today(), s.start_time)
                end = datetime.datetime.combine(datetime.date.today(), s.end_time)
                if end <= start:
                    end += datetime.timedelta(days=1)
                hours = (end - start).total_seconds() / 3600
                summary[uid]["total_hours"] += hours
                summary[uid]["chart_breakdown"][chart_name]["total_hours"] += hours
                
                # Count by shift name under the chart
                summary[uid]["chart_breakdown"][chart_name]["shifts"].setdefault(s.name, 0)
                summary[uid]["chart_breakdown"][chart_name]["shifts"][s.name] += 1
            else:
                # If no schedule, still count the chart occurrence
                summary[uid]["chart_breakdown"][chart_name]["shifts"].setdefault("No Shift", 0)
                summary[uid]["chart_breakdown"][chart_name]["shifts"]["No Shift"] += 1
        
        # Format results
        for uid in summary:
            summary[uid]["dates"].sort(key=lambda x: x["date"])
        
        # Sort by total duties desc
        result = sorted(summary.values(), key=lambda x: x['total_duties'], reverse=True)
        return Response(result)


class OfficeAdoptionReportView(APIView):
    permission_classes = [IsAdminOrSelf]

    def get(self, request):
        from org.models import WorkingOffice
        from django.db.models import Q, Prefetch
        from duties.models import DutyChart, Duty
        import datetime

        directorate_id = request.GET.get("directorate_id")
        ac_office_id = request.GET.get("ac_office_id")
        office_id = request.GET.get("office_id")

        # Prefetch with select_related and nested prefetch_related to load all data in bulk (4 queries total)
        offices = WorkingOffice.objects.select_related(
            "directorate", "ac_office", "cc_office"
        )

        # Fetch all offices in-memory to build recursive directorate resolution
        all_offices = list(offices)
        offices_by_id = {o.id: o for o in all_offices}

        # Helper to recursively resolve directorate ID
        def resolve_directorate_id(o, visited=None):
            if visited is None:
                visited = set()
            if o.id in visited:
                return None
            visited.add(o.id)
            if o.directorate_id:
                return o.directorate_id
            if o.ac_office and o.ac_office.directorate_id:
                return o.ac_office.directorate_id
            if o.parent_id:
                parent = offices_by_id.get(o.parent_id)
                if parent:
                    return resolve_directorate_id(parent, visited)
            return None

        # Filter working offices
        filtered_offices = all_offices
        if directorate_id and directorate_id != "all":
            target_dir_id = int(directorate_id)
            filtered_offices = [o for o in filtered_offices if resolve_directorate_id(o) == target_dir_id]
        if ac_office_id and ac_office_id != "all":
            filtered_offices = [o for o in filtered_offices if o.ac_office_id == int(ac_office_id)]
        if office_id and office_id != "all":
            filtered_offices = [o for o in filtered_offices if o.id == int(office_id)]

        filtered_office_ids = [o.id for o in filtered_offices]

        # Aggregate counts efficiently for the filtered offices
        from django.db.models import Count, Max
        stats_qs = WorkingOffice.objects.filter(id__in=filtered_office_ids).annotate(
            annotated_duty_count=Count('duties', distinct=True),
            annotated_chart_count=Count('duty_charts', distinct=True),
            max_duty_date=Max('duties__date'),
            max_chart_edited=Max('duty_charts__edited_at')
        )
        stats_map = {s.id: s for s in stats_qs}

        # Fetch duty charts and their associated user information efficiently
        charts_qs = DutyChart.objects.filter(office_id__in=filtered_office_ids).prefetch_related(
            Prefetch(
                "duties",
                queryset=Duty.objects.filter(user__isnull=False).select_related("user").only(
                    "id", "duty_chart_id", "user_id", "user__employee_id", "user__full_name"
                )
            )
        )
        
        charts_by_office = {}
        for c in charts_qs:
            charts_by_office.setdefault(c.office_id, []).append(c)

        # Build list of results
        office_list = []
        for office in filtered_offices:
            stats = stats_map.get(office.id)
            office_charts = charts_by_office.get(office.id, [])
            
            duty_chart_count = stats.annotated_chart_count if stats else 0
            duty_count = stats.annotated_duty_count if stats else 0
            has_started = duty_chart_count > 0 or duty_count > 0

            # Last activity
            last_activity = None
            if stats:
                max_edit = stats.max_chart_edited
                max_duty = stats.max_duty_date
                
                if max_edit:
                    last_activity = max_edit
                    
                if max_duty:
                    duty_dt = datetime.datetime.combine(max_duty, datetime.time.min)
                    if last_activity and timezone.is_aware(last_activity):
                        duty_dt = timezone.make_aware(duty_dt)
                    if not last_activity or duty_dt > last_activity:
                        last_activity = duty_dt

            charts_list = []
            for chart in office_charts:
                # Compute distinct emp_count and assigned_users entirely in-memory using prefetched objects
                chart_duties = [d for d in chart.duties.all() if d.user is not None]
                emp_ids = {d.user_id for d in chart_duties}
                emp_count = len(emp_ids)
                assigned_users_raw = {
                    (d.user.employee_id or "N/A", d.user.full_name or "Unknown")
                    for d in chart_duties
                }
                assigned_users = sorted(
                    [{"employee_id": uid, "name": name} for uid, name in assigned_users_raw],
                    key=lambda x: x["name"]
                )
                
                # Convert effective_date and end_date to Nepali BS dates
                try:
                    nep_start = nepali_datetime.date.from_datetime_date(chart.effective_date)
                    nepali_start_date = nep_start.strftime("%Y/%m/%d")
                except Exception:
                    nepali_start_date = chart.effective_date.isoformat()
                    
                try:
                    if chart.end_date:
                        nep_end = nepali_datetime.date.from_datetime_date(chart.end_date)
                        nepali_end_date = nep_end.strftime("%Y/%m/%d")
                    else:
                        nepali_end_date = nepali_start_date
                except Exception:
                    nepali_end_date = chart.end_date.isoformat() if chart.end_date else chart.effective_date.isoformat()

                charts_list.append({
                    "id": chart.id,
                    "name": chart.name or f"Roster Chart #{chart.id}",
                    "start_date": chart.effective_date.isoformat(),
                    "end_date": chart.end_date.isoformat() if chart.end_date else chart.effective_date.isoformat(),
                    "nepali_start_date": nepali_start_date,
                    "nepali_end_date": nepali_end_date,
                    "employee_count": emp_count,
                    "employees": list(assigned_users)
                })

            office_list.append({
                "id": office.id,
                "name": office.name,
                "directorate_name": office.directorate.directorate if office.directorate else "None",
                "ac_office_name": office.ac_office.name if office.ac_office else "None",
                "cc_office_name": office.cc_office.name if office.cc_office else "None",
                "duty_chart_count": duty_chart_count,
                "duty_count": duty_count,
                "last_activity": last_activity.isoformat() if last_activity else None,
                "has_started": has_started,
                "charts": charts_list
            })

        # Calculate summaries
        total_offices = len(office_list)
        started_offices = sum(1 for o in office_list if o["has_started"])
        not_started_offices = total_offices - started_offices
        adoption_rate = (started_offices / total_offices * 100) if total_offices > 0 else 0.0

        return Response({
            "summary": {
                "total_offices": total_offices,
                "started_offices": started_offices,
                "not_started_offices": not_started_offices,
                "adoption_rate": round(adoption_rate, 2)
            },
            "offices": office_list
        })